# --- Importações Essenciais ---
import streamlit as st
import pandas as pd
import os
import uuid
import shutil
import io
import io
import datetime
import docx
import re
import matplotlib.pyplot as plt
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from contextlib import redirect_stdout
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_community.llms import Ollama
from utils import (
    get_ollama_models,
    get_gemini_models,
    parse_agent_thoughts,
    display_formatted_thoughts,
    get_data_profile
)

def clean_markdown(text):
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'^#+\s', '', text, flags=re.MULTILINE)
    return text

def generate_docx_report(pinned_items, data_profile, user_name):
    doc = docx.Document()

    # Adiciona e centraliza o logotipo no início
    logo_path = "asset/LOGO.png"
    if os.path.exists(logo_path):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_logo = p_logo.add_run()
        run_logo.add_picture(logo_path, width=Inches(1.5))

    # Estilo e título principal
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)
    title = doc.add_heading('Relatório de Análise de Dados - JEDI', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadados
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.add_run(f"Gerado por: {user_name}\n").italic = True
    p_meta.add_run(f"Data: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M:%S')} ").italic = True
    doc.add_paragraph("---")

    # --- Seção do Sumário Estático ---
    doc.add_heading('Sumário', level=1)
    doc.add_paragraph("Perfil Inicial dos Dados", style='List Bullet')
    for i, item in enumerate(pinned_items):
        resumo_pergunta = item['user_prompt'][:60] + '...' if len(item['user_prompt']) > 60 else item['user_prompt']
        doc.add_paragraph(f"Interação {i+1}: {resumo_pergunta}", style='List Bullet')
    
    doc.add_page_break()

    # --- Seção do Perfil dos Dados ---
    if data_profile:
        doc.add_heading('Perfil Inicial dos Dados', level=1)
        
        def parse_and_add_profile(profile_string):
            lines = profile_string.strip().split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                if line.startswith('###'):
                    doc.add_heading(line.replace('###', '').strip(), level=2)
                elif line.startswith('- **'):
                    clean_line = line.replace('- **', '').replace('**', '').strip()
                    p = doc.add_paragraph()
                    p.add_run(clean_line).bold = True
                elif line.startswith('  - '):
                    clean_line = line.replace('`', '').strip()[2:] # Remove o marcador e os backticks
                    doc.add_paragraph(clean_line, style='List Bullet 2')
                else:
                    doc.add_paragraph(line)
        
        parse_and_add_profile(data_profile)
        doc.add_page_break()

    # --- Conteúdo Principal do Relatório ---
    for i, item in enumerate(pinned_items):
        heading_interaction = doc.add_heading(f"Interação {i+1}: Pergunta do Usuário", level=1)
        if i > 0:
            heading_interaction.paragraph_format.space_before = Pt(24)

        p_question = doc.add_paragraph()
        p_question.add_run(item['user_prompt']).italic = True
        
        heading_answer = doc.add_heading("Resposta do Agente", level=2)
        heading_answer.paragraph_format.space_before = Pt(18)

        content = clean_markdown(item['content'])
        p_answer = doc.add_paragraph(content)
        p_answer.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if item.get("image") and os.path.exists(item["image"]):
            p_image = doc.add_paragraph()
            p_image.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_image.paragraph_format.space_before = Pt(12)
            run = p_image.add_run()
            try:
                run.add_picture(item["image"], width=Inches(5.5))
            except Exception as e:
                doc.add_paragraph(f"(Erro ao adicionar imagem: {e})")
            p_caption = doc.add_paragraph()
            p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_caption.add_run(f"Figura {i+1}: Gráfico gerado pela análise.").italic = True

        if item.get("thoughts"):
            # O código para exibir os pensamentos permanece o mesmo
            pass

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio
def main_app():
    st.set_page_config(page_title="JEDI EDA", layout="wide")
    plots_dir = "temp_plots"

    def _pin_item(message):
        item_id = f"message-{message['timestamp']}-{hash(message['content'])}"
        for i, pinned_item in enumerate(st.session_state.pinned_items):
            if pinned_item.get("id") == item_id:
                st.session_state.pinned_items.pop(i)
                st.toast("Item despinado do relatório!", icon="📌")
                return
        user_prompt = ""
        for msg in reversed(st.session_state.messages):
            if msg['role'] == 'user':
                user_prompt = msg['content']
                break
        pinned_data = {
            "id": item_id, "role": "assistant", "content": message['content'],
            "timestamp": message['timestamp'], "image": message.get("image"),
            "thoughts": message.get("thoughts"), "user_prompt": user_prompt
        }
        st.session_state.pinned_items.append(pinned_data)
        st.toast("Item pinado para o relatório!", icon="✅")

    @st.cache_data
    def load_models():
        return get_ollama_models(), get_gemini_models()
    
    ollama_models, gemini_models = load_models()

    with st.sidebar:
        logo_path = "asset/LOGO.png"
        if os.path.exists(logo_path):
            st.image(logo_path, width=100)
        st.header(f"Bem-vindo, {st.session_state.get('user_name', 'Usuário')}!")
        if st.button("Logout", use_container_width=True):
            if os.path.exists(plots_dir):
                shutil.rmtree(plots_dir)
            st.session_state.clear()
            if "GOOGLE_API_KEY" in os.environ:
                del os.environ["GOOGLE_API_KEY"]
            st.rerun()

        if st.button("Reiniciar Conversa", use_container_width=True):
            st.session_state.messages = []
            st.session_state.pinned_items = []

            # Deletar a chave do arquivo atual força o bloco de recarregamento a ser executado
            if "current_file" in st.session_state:
                del st.session_state.current_file

            # Limpa os plots temporários
            if os.path.exists(plots_dir):
                shutil.rmtree(plots_dir)
            os.makedirs(plots_dir, exist_ok=True)
            
            st.toast("A conversa foi reiniciada. O arquivo precisa ser recarregado.", icon="🔄")
            st.rerun()
        st.divider()
        st.header("⚙️ Configurações")
        llm_provider = st.selectbox("Escolha o Provedor de LLM:", ["Ollama", "Gemini"], index=1)
        selected_model = None
        if llm_provider == "Ollama":
            if ollama_models:
                selected_model = st.selectbox("Escolha o Modelo Ollama:", ollama_models)
            else:
                st.warning("Nenhum modelo Ollama encontrado.")
        elif llm_provider == "Gemini":
            if gemini_models:
                filtered_gemini_models = [m for m in gemini_models if 'gemini' in m]
                default_model = 'models/gemini-2.0-flash'
                try:
                    default_index = filtered_gemini_models.index(default_model)
                except ValueError:
                    default_index = 0
                selected_model = st.selectbox("Escolha o Modelo Gemini:", filtered_gemini_models, index=default_index)
            else:
                st.warning("Nenhum modelo Gemini encontrado.")
        uploaded_file = st.file_uploader("Faça upload do seu arquivo CSV", type=["csv"])
        st.divider()
        st.header("Auditoria do Conselho")
        show_thoughts = st.toggle("Mostrar Diário de Bordo do Conselho", value=True, key="show_thoughts_toggle")
        st.divider()
        st.header("📌 Relatório Gerado")
        if not st.session_state.pinned_items:
            st.info("Nenhum item foi adicionado ao relatório ainda.")
        else:
            for item in st.session_state.pinned_items:
                with st.expander(f"Interação de {item['timestamp']}"):
                    with st.chat_message("user"):
                        st.markdown(item['user_prompt'])
                    with st.chat_message("assistant"):
                        st.markdown(item['content'])
                        if item.get("image") and os.path.exists(item["image"]):
                            image_path = item["image"]
                            if image_path.endswith(".html"):
                                with open(image_path, 'r', encoding='utf-8') as f:
                                    html_content = f.read()
                                st.components.v1.html(html_content, height=400, scrolling=True)
                            elif image_path.endswith(('.png', '.jpg', '.jpeg')):
                                st.image(image_path)
            if st.button("Limpar Itens Pinados", use_container_width=True):
                st.session_state.pinned_items = []
                st.toast("Itens pinados limpos!", icon="🧹")
                st.rerun()
            docx_data = generate_docx_report(st.session_state.pinned_items, st.session_state.data_profile, st.session_state.get('user_name', 'Usuário'))
            st.download_button(
                label="Download Relatório (.docx)",
                data=docx_data,
                file_name=f"relatorio_jedi_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

    st.title("✨ JEDI: Conselho de Análise de Dados")
    st.write("Converse com um Conselho de agentes Jedi para explorar seus dados, gerar insights e obter respostas inteligentes.")
    
    if uploaded_file is not None:
        if st.session_state.get("current_file") != uploaded_file.name:
            if os.path.exists(plots_dir): shutil.rmtree(plots_dir)
            os.makedirs(plots_dir, exist_ok=True)
            st.session_state.messages = []
            st.session_state.current_file = uploaded_file.name
            st.session_state.data_profile = None

        try:
            df = pd.read_csv(uploaded_file)
            st.success("Arquivo CSV carregado com sucesso!")
            st.dataframe(df.head())

            if st.session_state.data_profile is None:
                with st.spinner("Analisando o perfil dos dados..."):
                    st.session_state.data_profile = get_data_profile(df)
            
            with st.expander("Ver Perfil Detalhado dos Dados"):
                st.markdown(st.session_state.data_profile)

            llm = None
            if llm_provider == "Ollama":
                if selected_model: llm = Ollama(model=selected_model, temperature=0)
            elif llm_provider == "Gemini":
                if selected_model and os.getenv("GOOGLE_API_KEY"):
                    llm = ChatGoogleGenerativeAI(model=selected_model.replace('models/', ''), temperature=0)

            if "messages" not in st.session_state: st.session_state.messages = []

            for i, message in enumerate(st.session_state.messages):
                with st.chat_message(message["role"]):
                    st.markdown(message["content"])
                    if message["role"] == "assistant":
                        # Display plot if it exists
                        if "image" in message and os.path.exists(message["image"]):
                            image_path = message["image"]
                            if image_path.endswith(".png"):
                                st.image(image_path)
                        
                        # Display thoughts if they exist
                        if "thoughts" in message and message["thoughts"]:
                            with st.expander("Ver Diário de Bordo do Conselho Jedi 🧠"):
                                display_formatted_thoughts(message["thoughts"])
                        
                        # Pining button
                        item_id = f"message-{message['timestamp']}-{hash(message['content'])}"
                        is_pinned = any(item.get("id") == item_id for item in st.session_state.pinned_items)
                        pin_label = "📌 Pinado" if is_pinned else "🧷 Pinar"
                        if st.button(pin_label, key=f"pin_message_{i}"):
                            _pin_item(message)
                            st.rerun()

            if prompt := st.chat_input(f"{st.session_state.get('user_name', 'Usuário')}: Pergunte ao JEDI sobre os dados..."):
                st.session_state.messages.append({"role": "user", "content": prompt, "timestamp": datetime.datetime.now().isoformat()})
                with st.chat_message("user"):
                    st.markdown(prompt)

                # Adiciona verificação para o Ollama antes de prosseguir
                if llm_provider == "Ollama" and not selected_model:
                    error_message = "Parece que você selecionou o Ollama, mas nenhum modelo está disponível ou em execução. Por favor, inicie o Ollama localmente ou mude o provedor para 'Gemini' na barra lateral."
                    st.session_state.messages.append({"role": "assistant", "content": error_message, "timestamp": datetime.datetime.now().isoformat()})
                    st.rerun()
                
                elif llm:
                    with st.spinner("O Conselho Jedi está deliberando..."):
                        try:
                            from agents.master import run_jedi_council

                            council_response = run_jedi_council(llm, df, prompt, record_thoughts=show_thoughts)
                            
                            response_text = council_response.get("text_answer", "Ocorreu um erro ao processar a resposta.")
                            image_path = council_response.get("artifact_path")
                            thoughts = council_response.get("thoughts", [])

                            assistant_message = {
                                "role": "assistant", 
                                "content": response_text, 
                                "timestamp": datetime.datetime.now().isoformat(),
                                "thoughts": thoughts
                            }
                            if image_path:
                                assistant_message["image"] = image_path
                            
                            st.session_state.messages.append(assistant_message)
                            st.rerun()
                        except Exception as e:
                            st.error(f"Ocorreu um erro inesperado com o Conselho Jedi: {e}")
        except Exception as e:
            st.error(f"Ocorreu um erro ao carregar o arquivo CSV: {e}")
    else:
        st.markdown("---")
        st.markdown("<h3 style='text-align: center;'>🚀 Comece sua Análise de Dados!</h3>", unsafe_allow_html=True)
        st.markdown("---")
        st.write("<p>Para iniciar, siga os passos na barra lateral:</p>\n        <ol>\n            <li><strong>Faça o upload do seu arquivo CSV.</strong></li>\n            <li>O modelo <strong>Gemini 2.0 Flash</strong> já está selecionado por padrão. Se quiser, você pode trocá-lo.</li>\n            <li>Use a caixa de chat abaixo para começar a fazer perguntas!</li>\n        </ol>\n        <p><strong>Dica de Relatório:</strong></p>\n        <ul>\n            <li>Clique no ícone &#129527; ao lado de uma resposta para adicioná-la ao seu relatório.</li>\n            <li>O relatório é montado na barra lateral à esquerda.</li>\n            <li>Você pode baixar o relatório completo em formato <strong>.docx</strong> usando o botão na barra lateral.</li>\n        </ul>\n        <p><strong>Modo Desenvolvedor:</strong></p>\n        <p>Ative a opção \"Mostrar pensamentos do agente\" na barra lateral para entender como a IA está trabalhando!</p>", unsafe_allow_html=True)
